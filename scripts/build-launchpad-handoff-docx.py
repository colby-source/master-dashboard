"""Build LAUNCHPAD-DEV-HANDOFF.docx + LAUNCHPAD-DEV-ONBOARDING.docx from markdown sources."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
TARGETS = [
    (ROOT / "docs" / "LAUNCHPAD-DEV-HANDOFF.md", ROOT / "docs" / "LAUNCHPAD-DEV-HANDOFF.docx"),
    (ROOT / "docs" / "LAUNCHPAD-DEV-ONBOARDING.md", ROOT / "docs" / "LAUNCHPAD-DEV-ONBOARDING.docx"),
]

TEAL = RGBColor(0x01, 0x6F, 0x74)
SLATE = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def render_inline(paragraph, text: str):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("[") and "](" in part:
            label = part[1 : part.index("]")]
            run = paragraph.add_run(label)
            run.font.color.rgb = TEAL
            run.underline = True
        else:
            paragraph.add_run(part)


def build(src: Path, out: Path):
    doc = Document()
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE

    for level, size in [(1, 22), (2, 16), (3, 13)]:
        s = styles[f"Heading {level}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = TEAL

    def add_code(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = SLATE

    table_rows: list[list[str]] = []
    in_table = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        cleaned = []
        for r in table_rows:
            cells = [c.strip() for c in r]
            if all(re.match(r"^[-:]+$", c) for c in cells if c):
                continue
            cleaned.append(cells)
        if not cleaned:
            table_rows = []
            in_table = False
            return
        cols = max(len(r) for r in cleaned)
        t = doc.add_table(rows=len(cleaned), cols=cols)
        t.style = "Light Grid Accent 1"
        for ri, r in enumerate(cleaned):
            for ci in range(cols):
                cell = t.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                text = r[ci] if ci < len(r) else ""
                render_inline(p, text)
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()
        table_rows = []
        in_table = False

    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code("\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        elif in_table:
            flush_table()

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("─" * 60).font.color.rgb = MUTED
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            render_inline(p, re.sub(r"^\d+\.\s", "", line))
        elif line.strip() == "":
            pass
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

        i += 1

    flush_table()
    if in_code and code_buf:
        add_code("\n".join(code_buf))

    doc.save(out)
    print(f"wrote {out}")


for src, out in TARGETS:
    build(src, out)
